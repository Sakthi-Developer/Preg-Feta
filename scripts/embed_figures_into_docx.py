#!/usr/bin/env python3
"""Embed the nine generated figures into FETA_PREGNet_Manuscript.docx.

For each "Figure N (placeholder)." Block-Quotation paragraph, insert the matching
PNG (width 6.2in, centred) immediately above it, then rewrite the paragraph as a
bold "Figure N." caption run followed by the existing descriptive text, with the
word "(placeholder)" removed.
"""
from __future__ import annotations

import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCX = os.path.join(ROOT, "FETA_PREGNet_Manuscript.docx")
FIG = os.path.join(ROOT, "figures")

FIG_FILES = {
    1: "fig1_framework_overview.png",
    2: "fig2_feta_architecture.png",
    3: "fig3_pregnet_architecture.png",
    4: "fig4_cohort_trajectories.png",
    5: "fig5_roc_pr.png",
    6: "fig6_cv_forest.png",
    7: "fig7_confusion_matrices.png",
    8: "fig8_feta_attention.png",
    9: "fig9_pregnet_explainability.png",
}

doc = Document(DOCX)

placeholder_re = re.compile(r"^Figure\s+(\d+)\s*\(placeholder\)\.\s*(.*)$", re.DOTALL)

embedded = 0
for para in list(doc.paragraphs):
    m = placeholder_re.match(para.text.strip())
    if not m:
        continue
    n = int(m.group(1))
    rest = m.group(2).strip()
    png = os.path.join(FIG, FIG_FILES[n])
    if not os.path.exists(png):
        raise FileNotFoundError(png)

    # 1) Insert an image paragraph immediately above the caption paragraph.
    img_para = para.insert_paragraph_before()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    run.add_picture(png, width=Inches(6.2))

    # 2) Rewrite the caption paragraph: drop every existing run, then add a bold
    #    "Figure N." run and a normal run with the descriptive text.
    for r in list(para.runs):
        r._element.getparent().remove(r._element)
    cap = para.add_run(f"Figure {n}.")
    cap.bold = True
    body = para.add_run(f" {rest}")
    body.bold = False

    embedded += 1
    print(f"embedded Figure {n}: {FIG_FILES[n]}")

doc.save(DOCX)
print(f"\nSaved {DOCX} with {embedded} figures embedded.")
